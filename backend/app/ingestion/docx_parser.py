"""Word document extraction with python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.database.models import ExtractedBlock, ExtractedDocument
from app.ingestion.base_parser import DocumentParser
from app.logging_config import get_logger

logger = get_logger(__name__)


class DocxParseError(Exception):
    """Raised for malformed DOCX files."""


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    return style_name.lower().startswith("heading")


def _table_to_text(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        cells = [c for c in cells if c]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


class DocxParser(DocumentParser):
    extensions = {".docx"}

    def parse(self, path: Path, relative_path: str) -> ExtractedDocument:
        try:
            document = Document(str(path))
        except Exception as exc:  # noqa: BLE001
            raise DocxParseError(f"Cannot open DOCX: {exc}") from exc

        blocks: list[ExtractedBlock] = []
        current_heading: str | None = None

        # Iterate body elements in order
        try:
            body = document.element.body
        except Exception as exc:  # noqa: BLE001
            raise DocxParseError(f"Malformed DOCX body: {exc}") from exc

        from docx.oxml.ns import qn

        for child in body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                if not text:
                    continue
                if _is_heading(paragraph):
                    current_heading = text
                    blocks.append(
                        ExtractedBlock(
                            text=text,
                            heading=current_heading,
                            block_type="heading",
                        )
                    )
                else:
                    blocks.append(
                        ExtractedBlock(
                            text=text,
                            heading=current_heading,
                            block_type="paragraph",
                        )
                    )
            elif tag == qn("w:tbl"):
                table = Table(child, document)
                table_text = _table_to_text(table)
                if table_text.strip():
                    blocks.append(
                        ExtractedBlock(
                            text=table_text,
                            heading=current_heading,
                            block_type="table",
                        )
                    )

        if not blocks:
            # Fallback: paragraphs only
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    blocks.append(ExtractedBlock(text=text, block_type="paragraph"))

        if not blocks:
            raise DocxParseError("Empty DOCX (no extractable text)")

        return ExtractedDocument(
            filepath=relative_path.replace("\\", "/"),
            filename=path.name,
            file_type="docx",
            blocks=blocks,
            metadata={},
        )
