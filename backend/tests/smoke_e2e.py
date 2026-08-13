"""End-to-end smoke test against a tiny fixture corpus."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

BACKEND = Path(__file__).resolve().parents[1]
ROOT = BACKEND.parent
sys.path.insert(0, str(BACKEND))

import pymupdf as fitz
from docx import Document

from app.config import Settings
from app.database.models import ChatMode
from app.database.qdrant import QdrantStore
from app.database.sqlite import SQLiteDatabase
from app.llm.answer_generator import AnswerGenerator
from app.llm.ollama_client import OllamaClient
from app.retrieval.embeddings import EmbeddingService
from app.retrieval.hybrid_search import HybridSearch
from app.retrieval.keyword_search import KeywordSearch
from app.retrieval.vector_search import VectorSearch
from app.services.chat_service import ChatService
from app.services.ingestion_service import IngestionService
from app.services.search_service import SearchService


def main() -> int:
    docs = ROOT / "data" / "smoke_docs"
    if docs.exists():
        shutil.rmtree(docs)
    pdf_dir = docs / "Year4" / "Cloud_Security" / "Assignments"
    pdf_dir.mkdir(parents=True)
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "NIS2 compliance and Docker containers on AWS EC2 for cloud security.",
    )
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Active Directory and Linux hardening notes.")
    doc.save(pdf_dir / "Cloud_CA2.pdf")
    doc.close()

    docx_dir = docs / "Year3" / "Network_Security"
    docx_dir.mkdir(parents=True)
    w = Document()
    w.add_heading("Firewall Configuration", 1)
    w.add_paragraph(
        "Docker was selected for lab isolation and packet capture with Wireshark."
    )
    w.save(docx_dir / "Lab_Notes.docx")

    sqlite_path = ROOT / "data" / "sqlite" / "smoke.db"
    qdrant_path = ROOT / "data" / "qdrant_smoke"
    if sqlite_path.exists():
        sqlite_path.unlink()
    if qdrant_path.exists():
        shutil.rmtree(qdrant_path)

    settings = Settings(
        documents_path=str(docs.resolve()),
        sqlite_path=str(sqlite_path),
        qdrant_path=str(qdrant_path),
        qdrant_collection="smoke_chunks",
    )
    db = SQLiteDatabase(settings.resolve_sqlite_path())
    qdrant = QdrantStore(settings.resolve_qdrant_path(), settings.qdrant_collection)
    ollama = OllamaClient(
        settings.ollama_base_url,
        settings.ollama_chat_model,
        settings.ollama_embed_model,
    )
    embeddings = EmbeddingService(ollama)
    ingestion = IngestionService(settings, db, qdrant, embeddings)

    stats1 = ingestion.run()
    assert stats1.processed == 2, stats1
    stats2 = ingestion.run()
    assert stats2.skipped == 2, stats2

    hits = db.keyword_search("NIS2")
    assert hits and "NIS2" in hits[0][0].text

    hybrid = HybridSearch(VectorSearch(qdrant, db, embeddings), KeywordSearch(db))
    search = SearchService(settings, db, hybrid)
    chat = ChatService(settings, search, AnswerGenerator(ollama))
    result = chat.chat("What did I use Docker for?", mode=ChatMode.ASK)
    assert result.answer
    assert result.sources
    print("SMOKE OK")
    print(result.answer[:500])
    print("sources:", [(s.filename, s.page, s.heading) for s in result.sources])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
