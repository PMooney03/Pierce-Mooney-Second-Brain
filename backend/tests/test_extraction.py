"""PDF and DOCX extraction tests with generated fixtures."""

from __future__ import annotations

from pathlib import Path

import pymupdf as fitz
from docx import Document

from app.ingestion.docx_parser import DocxParser
from app.ingestion.pdf_parser import PdfParser


def _make_pdf(path: Path, pages: list[str]) -> None:
    doc = fitz.open()
    for text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Network Security", level=1)
    doc.add_paragraph("Firewalls filter traffic between trust zones.")
    doc.add_heading("Implementation", level=2)
    doc.add_paragraph("Docker was selected for lab isolation.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tool"
    table.cell(0, 1).text = "Use"
    table.cell(1, 0).text = "Wireshark"
    table.cell(1, 1).text = "Packet capture"
    doc.save(path)


def test_pdf_extraction_preserves_pages(tmp_path: Path) -> None:
    pdf = tmp_path / "lecture.pdf"
    _make_pdf(pdf, ["Active Directory basics", "Group Policy objects"])
    extracted = PdfParser().parse(pdf, "Year2/Sec/lecture.pdf")
    assert extracted.file_type == "pdf"
    assert extracted.filename == "lecture.pdf"
    pages = {b.page for b in extracted.blocks}
    assert 1 in pages and 2 in pages
    joined = " ".join(b.text for b in extracted.blocks)
    assert "Active Directory" in joined
    assert "Group Policy" in joined


def test_docx_extraction_preserves_headings_and_order(tmp_path: Path) -> None:
    path = tmp_path / "notes.docx"
    _make_docx(path)
    extracted = DocxParser().parse(path, "Year3/Net/notes.docx")
    assert extracted.file_type == "docx"
    types = [b.block_type for b in extracted.blocks]
    assert "heading" in types
    assert "paragraph" in types
    assert "table" in types
    texts = [b.text for b in extracted.blocks]
    assert any("Network Security" in t for t in texts)
    assert any("Docker" in t for t in texts)
    assert any("Wireshark" in t for t in texts)
    # Heading metadata on following paragraph
    docker_blocks = [b for b in extracted.blocks if "Docker" in b.text]
    assert docker_blocks
    assert docker_blocks[0].heading == "Implementation"
